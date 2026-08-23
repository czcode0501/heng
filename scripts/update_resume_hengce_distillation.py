from pathlib import Path

from docx import Document


SOURCE = Path(r"D:\桌面归档\文件\简历\刚晨钊_2026届简历_衡策QuantDesk_投资经理版.docx")
OUTPUT = Path(r"D:\桌面归档\文件\简历\刚晨钊_2026届简历_衡策QuantDesk_三Skill蒸馏版.docx")

REPLACEMENTS = {
    "独立规划并落地面向中美市场的量化交易平台": (
        "独立设计并开发面向中美市场的可切换投资经理量化平台，以多源方法论/工作流蒸馏技术整合 "
        "Augur、InvestorSkills 与 AI Berkshire 的因子归因、完整投资流程和证据核验机制，构建“筛选—研究—反方挑战—决策—监控”统一契约；"
        "将巴菲特等7位投资大师的公开理念配置为可切换基金经理，并把因子贡献、双源财务核验、超过1%的数据冲突闸门及投资论文复核分别落地于经理面板、个股分析和持仓评价"
    ),
    "围绕开源增长制定用户分层运营方案": (
        "围绕产品开源与用户增长制定分层运营方案：规划游客免注册体验、注册用户专属功能与基金经理策略库扩展，设计“反馈收集—需求分级—版本验证—指标复盘”闭环；"
        "以用户激活率、游客转注册率、核心功能使用率和留存率评估运营效果，并据此推动产品功能迭代与落地"
    ),
}


def replace_paragraph_text(paragraph, value: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(value)
        return
    paragraph.runs[0].text = value
    for run in paragraph.runs[1:]:
        run.text = ""


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    document = Document(SOURCE)
    replaced = set()
    for paragraph in document.paragraphs:
        for prefix, value in REPLACEMENTS.items():
            if paragraph.text.startswith(prefix):
                replace_paragraph_text(paragraph, value)
                replaced.add(prefix)
    missing = set(REPLACEMENTS) - replaced
    if missing:
        raise RuntimeError(f"未找到待替换段落: {sorted(missing)}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
